from docx import Document
import re

def add_question_numbers_docx(input_file, output_file, question_type):
    # 打开输入文件
    doc = Document(input_file)

    question_type = int(question_type)

    # 提取文本内容
    input_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    if not input_text.strip():
        print("输入文件为空或没有有效内容！")
        return

    # 匹配题目块的正则表达式
    # 根据 question_type 设置正则表达式
    if question_type == 1:  # 选择题
        pattern = r"(.*?\nA\..*?\nB\..*?\nC\..*?\nD\..*?\n答案：.*?)(?=\n.*?\nA\.|\Z)"
    elif question_type == 2:  # 判断题
        # pattern = r"(.*?)(?:\n|$)答案：(正确|错误)"
        # pattern = r"(.*?)(?:\n答案：(正确|错误))"
        pattern = r"([^\n]+)(?:\n答案：(正确|错误))"  # 匹配一行题目后跟答案部分
    matches = re.findall(pattern, input_text, re.S)

    if not matches:
        print("未找到任何匹配的题目格式，请检查输入内容是否符合格式要求！")
        return
    
    total = 0

    # 为每个题目块添加序号
    numbered_questions = []
    for i, match in enumerate(matches, start=1):
        if question_type == 1:
            # 去掉题目块内多余的换行符
            cleaned_match = re.sub(r"\n+", "\n", match.strip())
            lines = cleaned_match.split("\n")
            lines[0] = f"{i}. {lines[0]}"  # 为题目第一行添加序号
            numbered_questions.append("\n".join(lines))
        elif question_type == 2:
            # match 是一个元组 (题目, 答案)
            question, answer = match
            question = question.strip()
            # 去除题目开头的多余符号（如句号、空格等）
            question = re.sub(r"^[^\w]+", "", question)  # 去除题目开头的非单词字符（例如：句号、空格）
            # 为判断题添加序号
            question_with_number = f"{i}. {question}"
            # 格式化输出
            numbered_questions.append(f"{question_with_number}\n答案：{answer.strip()}")
        
        total = i

    # 写入到输出文件
    output_doc = Document()
    for question in numbered_questions:
        for line in question.split("\n"):
            output_doc.add_paragraph(line)
        output_doc.add_paragraph("")  # 添加空行分隔题目

    output_doc.save(output_file)
    print(f"已处理完成，共 {total} 道题")